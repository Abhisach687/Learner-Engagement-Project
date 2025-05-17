#code to extract tables from a DOCX file, convert them to PNG images, and save them in a new DOCX file

import os
import re
import mammoth
import imgkit
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph as DocxParagraph

# Configuration
# DOCX_PATH = r"C:\Users\abhis\Downloads\Documents\Thesis report images\Report_draft_comp.docx"
# OUTPUT_DIR = r"C:\Users\abhis\Downloads\Documents\Thesis report images\table_png"
DOCX_PATH = r"C:\Users\abhis\Downloads\Documents\Thesis report images\AppendixA.docx"
OUTPUT_DIR = r"C:\Users\abhis\Downloads\Documents\Thesis report images"
IMGKIT_OPTIONS = {
    "format": "png",
    "encoding": "UTF-8",
    "quality": "100",
}

# Helpers
def clean_filename(s, maxlen=50):
    return re.sub(r"[^0-9A-Za-z]+", "_", s).strip("_")[:maxlen]

# Extract subheadings above each table
def extract_headings(docx_path):
    doc = DocxDocument(docx_path)
    headings = []
    last = None
    for block in doc.element.body.iterchildren():
        if isinstance(block, CT_P):
            text = DocxParagraph(block, doc).text.strip()
            if text:
                last = text
        elif isinstance(block, CT_Tbl):
            headings.append(last or f"table_{len(headings)+1}")
            last = None
    return headings

# Function to wrap raw table HTML with styling
def table_to_html(raw_html):
    css = """
    <style>
      table {
        border-collapse: collapse;
        font-family: Arial, sans-serif;
        width: 100%;
      }
      th, td {
        border: 1px solid #999999;
        padding: 10px 8px;
        font-size: 11pt;
        vertical-align: middle;
        white-space: normal;
        overflow-wrap: break-word;
        background: #FFFFFF;
      }
      th {
        background-color: #F1F3F5;
        text-align: left;
        font-weight: normal;
      }
      tr:nth-child(odd) td {
        background-color: #FFFFFF;
      }
      tr:nth-child(even) td {
        background-color: #F9FAFB;
      }
    </style>
    """
    html = f"<html><head>{css}</head><body>{raw_html}</body></html>"
    return html

# 1) Convert DOCX to HTML via Mammoth
with open(DOCX_PATH, 'rb') as docx_file:
    result = mammoth.convert_to_html(docx_file)
html_output = result.value
soup = BeautifulSoup(html_output, 'html.parser')
tables = soup.find_all('table')
headings = extract_headings(DOCX_PATH)

# Prepare output directory
os.makedirs(OUTPUT_DIR, exist_ok=True)
png_paths = []

# 2) Render each table into a PNG via imgkit
for idx, tbl in enumerate(tables, start=1):
    raw = str(tbl)
    wrapped = table_to_html(raw)
    heading = headings[idx-1] if idx-1 < len(headings) else f"table_{idx:02d}"
    fname = f"{idx:02d}_{clean_filename(heading)}.png"
    out_path = os.path.join(OUTPUT_DIR, fname)
    imgkit.from_string(wrapped, out_path, options=IMGKIT_OPTIONS)
    png_paths.append(out_path)
    print(f"Rendered table {idx:02d} -> {fname}")

# 3) Rebuild a new DOCX with images replacing tables
orig = DocxDocument(DOCX_PATH)
new = DocxDocument()
# Copy page settings
src_sec, dst_sec = orig.sections[0], new.sections[0]
for prop in ('page_width','page_height','left_margin','right_margin','top_margin','bottom_margin'):
    setattr(dst_sec, prop, getattr(src_sec, prop))

img_iter = iter(png_paths)
for block in orig.element.body.iterchildren():
    if isinstance(block, CT_P):
        new.add_paragraph(DocxParagraph(block, orig).text)
    elif isinstance(block, CT_Tbl):
        try:
            img_file = next(img_iter)
            width_emu = src_sec.page_width - src_sec.left_margin - src_sec.right_margin
            # Convert EMU to inches if attribute exists
            if hasattr(width_emu, 'inches'):
                width_in = width_emu.inches
            else:
                width_in = width_emu / 914400.0
            new.add_picture(img_file, width=Inches(width_in))
        except StopIteration:
            break

# 4) Save the reconstructed document
out_docx = os.path.join(OUTPUT_DIR, 'Report_with_images.docx')
new.save(out_docx)
print(f"Done! New document saved to: {out_docx}")

# code to replace tables in a DOCX file with images from a specified directory
# and save the modified document to a new DOCX file

# python replace_tables_with_images.py `
#   "C:\Users\abhis\Downloads\Documents\Thesis report images\Report_draft_comp.docx" `
#   "C:\Users\abhis\Downloads\Documents\Thesis report images\table_png" `
#   "C:\Users\abhis\Downloads\Documents\Thesis report images\Report_with_images.docx"


# import os
# from docx import Document
# from docx.shared import Inches
# from docx.oxml import OxmlElement
# from docx.oxml.table import CT_Tbl
# from docx.oxml.text.paragraph import CT_P
# from docx.text.paragraph import Paragraph
# import re

# # === User-specific paths (no CLI args needed) ===
# DOCX_PATH = r"C:\Users\abhis\Downloads\Documents\Thesis report images\Report_draft_comp.docx"
# IMG_FOLDER = r"C:\Users\abhis\Downloads\Documents\Thesis report images\table_png"
# OUTPUT_DOCX = r"C:\Users\abhis\Downloads\Documents\Thesis report images\Report_with_images.docx"

# # === Helpers ===
# def clean_filename(s, maxlen=50):
#     return re.sub(r"[^0-9A-Za-z]+", "_", s).strip("_")[:maxlen]

# # Extract the last non-empty paragraph text before each table
# from docx import Document as _Docx

# def extract_headings(doc):
#     headings = []
#     last_text = None
#     for block in doc.element.body.iterchildren():
#         if isinstance(block, CT_P):
#             text = Paragraph(block, doc).text.strip()
#             if text:
#                 last_text = text
#         elif isinstance(block, CT_Tbl):
#             headings.append(last_text or "")
#             last_text = None
#     return headings

# # === Main replacement logic ===
# def replace_tables_with_images():
#     # Load document
#     doc = Document(DOCX_PATH)
#     sec = doc.sections[0]
#     usable_width_in = (sec.page_width - sec.left_margin - sec.right_margin) / 914400.0

#     # 0) Set default font to Arial
#     normal_style = doc.styles['Normal']
#     normal_style.font.name = 'Arial'
#     for style_name in ['Heading 1','Heading 2','Heading 3','Title','Subtitle']:
#         if style_name in doc.styles:
#             doc.styles[style_name].font.name = 'Arial'

#     # 1) Clean up manual line breaks in runs:
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             if '\n' in run.text:
#                 run.text = run.text.replace('\n', ' ')
#     # and inside table cells
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     for run in paragraph.runs:
#                         if '\n' in run.text:
#                             run.text = run.text.replace('\n', ' ')

#     # 2) Resize existing inline images if too wide:
#     for shape in doc.inline_shapes:
#         try:
#             # scale down only if shape exceeds text width
#             if shape.width > Inches(usable_width_in):
#                 shape.width = Inches(usable_width_in)
#         except Exception:
#             pass

#     # 3) Extract headings and map to image files
#     headings = extract_headings(doc)
#     img_files = sorted([f for f in os.listdir(IMG_FOLDER) if f.lower().endswith('.png')])
#     img_map = {}
#     for idx, heading in enumerate(headings, start=1):
#         prefix = f"{idx:02d}_{clean_filename(heading)}".lower()
#         match = next((f for f in img_files if f.lower().startswith(prefix)), None)
#         if not match and idx-1 < len(img_files):
#             match = img_files[idx-1]
#             print(f"⚠️ No exact image match for '{heading}', using '{match}'")
#         img_map[idx] = match

#     # 4) Replace each table in document order
#     for idx, table in enumerate(list(doc.tables), start=1):
#         img_name = img_map.get(idx)
#         if not img_name:
#             continue
#         img_path = os.path.join(IMG_FOLDER, img_name)
#         # Insert picture after table node
#         tbl_elm = table._element
#         pic_p = OxmlElement('w:p')
#         tbl_elm.addnext(pic_p)
#         para = Paragraph(pic_p, table._parent)
#         run = para.add_run()
#         run.add_picture(img_path, width=Inches(usable_width_in))
#         # Remove original table
#         tbl_elm.getparent().remove(tbl_elm)

#     # 5) Save output
#     doc.save(OUTPUT_DOCX)
#     print(f"✅ Tables replaced and formatting fixed. Output saved to: {OUTPUT_DOCX}")

# # Run
# replace_tables_with_images()
